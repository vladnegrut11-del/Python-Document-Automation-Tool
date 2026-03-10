import pandas as pd
from docx import Document
import os
import re
from pathlib import Path

def replace_text_in_document(doc, replacements):
    """
    Replace text in document paragraphs, tables, and headers/footers
    """
    # Sort replacements by key length (longest first) to avoid partial matches
    sorted_replacements = dict(sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True))
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        for key, value in sorted_replacements.items():
            if key in paragraph_text:
                # Replace text while preserving formatting
                runs = paragraph.runs
                # Combine all text from runs
                full_text = ''.join(run.text for run in runs)
                
                # Replace the key with value
                if key in full_text:
                    new_text = full_text.replace(key, str(value))
                    
                    # Clear existing runs
                    for run in runs:
                        run.text = ""
                    
                    # Put new text in first run
                    if runs:
                        runs[0].text = new_text
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in sorted_replacements.items():
                    cell_text = cell.text
                    if key in cell_text:
                        for paragraph in cell.paragraphs:
                            runs = paragraph.runs
                            full_text = ''.join(run.text for run in runs)
                            
                            if key in full_text:
                                new_text = full_text.replace(key, str(value))
                                
                                # Clear existing runs
                                for run in runs:
                                    run.text = ""
                                
                                # Put new text in first run
                                if runs:
                                    runs[0].text = new_text
    
    # Replace in headers and footers
    for section in doc.sections:
        # Header
        header = section.header
        for paragraph in header.paragraphs:
            for key, value in sorted_replacements.items():
                runs = paragraph.runs
                full_text = ''.join(run.text for run in runs)
                
                if key in full_text:
                    new_text = full_text.replace(key, str(value))
                    
                    # Clear existing runs
                    for run in runs:
                        run.text = ""
                    
                    # Put new text in first run
                    if runs:
                        runs[0].text = new_text
        
        # Footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            for key, value in sorted_replacements.items():
                runs = paragraph.runs
                full_text = ''.join(run.text for run in runs)
                
                if key in full_text:
                    new_text = full_text.replace(key, str(value))
                    
                    # Clear existing runs
                    for run in runs:
                        run.text = ""
                    
                    # Put new text in first run
                    if runs:
                        runs[0].text = new_text

def generate_auto_fields(index, replacements):
    """
    Generate automatic fields that can be used in templates
    """
    from datetime import datetime
    
    auto_fields = {}
    
    # Row number (1-based)
    auto_fields['ROW_NUMBER'] = str(index + 1)
    auto_fields['INDEX'] = str(index)
    
    # Current date in various formats
    now = datetime.now()
    auto_fields['DATA_CURENTA'] = now.strftime('%d.%m.%Y')
    auto_fields['DATA_CURENTA_LUNG'] = now.strftime('%d %B %Y')
    auto_fields['ANUL_CURENT'] = now.strftime('%Y')
    auto_fields['LUNA_CURENTA'] = now.strftime('%m')
    auto_fields['ZIUA_CURENTA'] = now.strftime('%d')
    
    # Full name combinations if NUME and PRENUME exist
    if 'NUME' in replacements and 'PRENUME' in replacements:
        nume = str(replacements['NUME']).strip() if not pd.isna(replacements['NUME']) else ""
        prenume = str(replacements['PRENUME']).strip() if not pd.isna(replacements['PRENUME']) else ""
        
        if nume and prenume:
            auto_fields['NUME_COMPLET'] = f"{nume} {prenume}"
            auto_fields['PRENUME_NUME'] = f"{prenume} {nume}"
            auto_fields['INITIALE'] = f"{prenume[0] if prenume else ''}.{nume[0] if nume else ''}."
    
    return auto_fields

def safe_replace_in_filename(text, replacements):
    """
    Safely replace placeholders in filename, avoiding partial matches using word boundaries
    """
    # Sort replacements by key length (longest first) to avoid partial matches
    sorted_replacements = dict(sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True))
    
    result = text
    for key, value in sorted_replacements.items():
        if key in result:
            # Use word boundaries to avoid partial replacements
            pattern = r'\b' + re.escape(key) + r'\b'
            replacement_value = str(value).strip() if not pd.isna(value) and str(value).strip() else ""
            result = re.sub(pattern, replacement_value, result)
    
    return result

def sanitize_filename(filename):
    """
    Remove or replace invalid characters for filenames
    """
    # Replace invalid characters with space
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, ' ')
    
    # Remove multiple consecutive spaces
    filename = re.sub(r' +', ' ', filename)
    
    # Remove leading/trailing spaces
    filename = filename.strip()
    
    return filename

def process_single_template(template_path, df, output_folder):
    """
    Process a single template with all Excel data
    """
    template_name = Path(template_path).stem
    print(f"\nProcessing template: {template_name}")
    
    # Create subfolder for this template
    template_folder = Path(output_folder) / sanitize_filename(template_name)
    template_folder.mkdir(parents=True, exist_ok=True)
    
    generated_count = 0
    
    # Process each row in the Excel file
    for index, row in df.iterrows():
        try:
            # Load template document with error handling
            try:
                doc = Document(template_path)
            except Exception as doc_error:
                print(f"  Error loading template: {doc_error}")
                print(f"  Try converting {template_name} to .docx format")
                return 0
            
            # Create replacements dictionary from row data
            replacements = {}
            for column in df.columns:
                value = row[column]
                # Handle NaN values
                if pd.isna(value):
                    value = ""
                replacements[column] = value
            
            # Add auto-generated fields
            auto_fields = generate_auto_fields(index, replacements)
            replacements.update(auto_fields)
            
            # Print auto-generated fields for first few rows (debugging)
            if index < 3:
                print(f"  Row {index + 1} custom fields:")
                for key, value in auto_fields.items():
                    if key in ['SERIE PVRTL', 'BULETIN PP', 'PVPIF', 'schita', 'imputernicire', 'ORDIN INCEPERE LUCRARI']:
                        print(f"    {key}: {value}")
            
            # Replace text in document
            replace_text_in_document(doc, replacements)
            
            # Generate filename based on template name with safe replacements
            filename = template_name
            
            # Use safe replacement function to avoid partial matches
            filename = safe_replace_in_filename(filename, replacements)
            
            # If no replacements were made or filename is empty, use fallback
            if filename == template_name or not filename.strip():
                if 'NUME' in df.columns and 'PRENUME' in df.columns:
                    nume = str(row['NUME']).strip() if not pd.isna(row['NUME']) else f"Row{index+1}"
                    prenume = str(row['PRENUME']).strip() if not pd.isna(row['PRENUME']) else ""
                    filename = f"{nume} {prenume} {template_name}".strip()
                else:
                    filename = f"{template_name}_Row_{index+1}"
            
            # Sanitize filename
            filename = sanitize_filename(filename)
            
            # Save document
            output_path = template_folder / f"{filename}.docx"
            doc.save(output_path)
            
            generated_count += 1
            
            if (index + 1) % 50 == 0:  # Progress update every 50 files
                print(f"  Progress: {index + 1} documents generated...")
            
        except Exception as e:
            print(f"  Error processing row {index + 1}: {e}")
            continue
    
    print(f"  Completed: {generated_count} documents generated in '{template_folder}'")
    return generated_count

def generate_documents_from_multiple_templates(template_paths, excel_path, output_folder="output"):
    """
    Generate personalized documents from multiple templates and Excel data
    
    Args:
        template_paths (list): List of paths to template documents
        excel_path (str): Path to the Excel file with data
        output_folder (str): Folder to save generated documents
    """
    
    # Create output folder if it doesn't exist
    Path(output_folder).mkdir(parents=True, exist_ok=True)
    
    # Read Excel file
    try:
        df = pd.read_excel(excel_path)
        print(f"Loaded Excel file with {len(df)} rows and columns: {list(df.columns)}")
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return
    
    total_generated = 0
    successful_templates = 0
    
    # Process each template
    for i, template_path in enumerate(template_paths, 1):
        print(f"\n{'='*50}")
        print(f"Template {i}/{len(template_paths)}: {Path(template_path).name}")
        print(f"{'='*50}")
        
        # Validate template exists
        if not os.path.exists(template_path):
            print(f"  Error: Template file not found: {template_path}")
            continue
            
        # Check if template is .docx format
        if not template_path.lower().endswith('.docx'):
            print(f"  Warning: Template '{Path(template_path).name}' is not .docx format")
            print(f"  Skipping... Please convert to .docx first")
            continue
        
        # Process this template
        generated = process_single_template(template_path, df, output_folder)
        if generated > 0:
            total_generated += generated
            successful_templates += 1
    
    print(f"\n{'='*60}")
    print(f"SUMMARY:")
    print(f"  Templates processed: {successful_templates}/{len(template_paths)}")
    print(f"  Total documents reported: {total_generated}")
    
    # Verify actual total count
    output_path = Path(output_folder)
    actual_total = len(list(output_path.rglob("*.docx")))
    print(f"  Actual documents found: {actual_total}")
    
    if actual_total != total_generated:
        print(f"  DISCREPANCY: {total_generated - actual_total} documents missing!")
        print(f"  Check the warnings above for failed saves.")
    
    print(f"  Output location: {Path(output_folder).absolute()}")
    print(f"{'='*60}")

def get_templates_from_folder(folder_path):
    """
    Get all .docx templates from a folder
    """
    folder = Path(folder_path)
    if not folder.exists():
        return []
    
    templates = []
    for file in folder.glob("*.docx"):
        if not file.name.startswith("~"):  # Skip temp files
            templates.append(str(file))
    
    return sorted(templates)

def main():
    """
    Funcția principală pentru generarea documentelor din multiple template-uri
    """
    print("Generator Documente din Multiple Template-uri")
    print("=" * 60)
    
    # Get Excel file path
    excel_path = input("Introdu calea către fișierul Excel (.xlsx sau .xls): ").strip().strip('"')
    
    if not os.path.exists(excel_path):
        print(f"Eroare: Fișierul Excel nu a fost găsit: {excel_path}")
        return
    
    # Get templates - multiple options
    print("\nChoose template input method:")
    print("1. Select individual template files")
    print("2. Select all templates from a folder")
    
    choice = input("Enter choice (1 or 2): ").strip()
    
    template_paths = []
    
    if choice == "2":
        # Get all templates from folder
        folder_path = input("Enter path to templates folder: ").strip().strip('"')
        template_paths = get_templates_from_folder(folder_path)
        
        if not template_paths:
            print(f"No .docx templates found in: {folder_path}")
            return
        
        print(f"\nFound {len(template_paths)} templates:")
        for i, template in enumerate(template_paths, 1):
            print(f"  {i}. {Path(template).name}")
        
        confirm = input(f"\nProcess all {len(template_paths)} templates? (y/n): ").lower().strip()
        if confirm != 'y':
            return
            
    else:
        # Get individual template files
        print("\nEnter template file paths (one per line, empty line to finish):")
        while True:
            template_path = input("Template path: ").strip().strip('"')
            if not template_path:
                break
            
            if os.path.exists(template_path) and template_path.lower().endswith('.docx'):
                template_paths.append(template_path)
                print(f"  Added: {Path(template_path).name}")
            else:
                print(f"  Error: File not found or not .docx: {template_path}")
        
        if not template_paths:
            print("No valid templates provided.")
            return
    
    # Get output folder
    output_folder = input("\nEnter output folder name (press Enter for 'output'): ").strip()
    if not output_folder:
        output_folder = "output"
    
    # Generate documents
    print(f"\nStarting generation with {len(template_paths)} templates...")
    generate_documents_from_multiple_templates(template_paths, excel_path, output_folder)

if __name__ == "__main__":
    # Example usage for multiple specific templates:
    # generate_documents_from_multiple_templates(
    #     template_paths=[
    #         "template1.docx",
    #         "template2.docx",
    #         "template3.docx"
    #     ],
    #     excel_path="data.xlsx",
    #     output_folder="all_generated_documents"
    # )
    
    # Example usage for all templates in a folder:
    # templates = get_templates_from_folder("templates_folder")
    # generate_documents_from_multiple_templates(templates, "data.xlsx", "output")
    
    # Run interactive mode
    main()